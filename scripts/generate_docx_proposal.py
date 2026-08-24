import os
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

def generate_academic_proposal():
    template_path = "/Users/yuanyi/Downloads/32e9bad9-cf02-4c40-b19c-c0043f3f5179.docx"
    out_path = "docs/qwen/千问办公专家套件方案书-多专家协同研发工作流.docx"
    download_copy = "/Users/yuanyi/Downloads/千问办公专家套件方案书-多专家协同研发工作流.docx"

    doc = docx.Document(template_path)

    # Standard Academic A4 Margins (GB/T 7713: 2.54cm top/bottom, 3.18cm left/right or 2.8cm)
    for section in doc.sections:
        section.top_margin = Inches(1.0)      # 2.54 cm
        section.bottom_margin = Inches(1.0)   # 2.54 cm
        section.left_margin = Inches(1.1)     # 2.8 cm
        section.right_margin = Inches(1.1)    # 2.8 cm

    # Academic Palette: Standard Black & Charcoal
    COLOR_BLACK = RGBColor(0, 0, 0)
    COLOR_HEADING = RGBColor(17, 24, 39)
    COLOR_MUTED = RGBColor(75, 85, 99)

    # Default Normal Style to Academic SongTi + Times New Roman
    for s_name in ["Normal", "正文"]:
        if s_name in doc.styles:
            s = doc.styles[s_name]
            s.font.name = "Times New Roman"
            s.font.size = Pt(10.5)
            rFonts = s.element.rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), "宋体")
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:cs"), "Times New Roman")

    def clear_paragraph(p):
        for r in list(p.runs):
            p._p.remove(r._r)

    def format_academic_run(run, font_size_pt, font_type="body", bold=False, color=COLOR_BLACK):
        run.font.size = Pt(font_size_pt)
        run.bold = bold
        run.font.color.rgb = color
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        if font_type in ["title", "heading"]:
            rFonts.set(qn("w:eastAsia"), "黑体")
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:cs"), "Times New Roman")
        elif font_type == "meta":
            rFonts.set(qn("w:eastAsia"), "楷体")
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:cs"), "Times New Roman")
        else:  # body / table
            rFonts.set(qn("w:eastAsia"), "宋体")
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:cs"), "Times New Roman")
        rFonts.set(qn("w:hint"), "eastAsia")

    def format_paragraph(p, space_before=3, space_after=3, line_spacing=1.3, align=WD_ALIGN_PARAGRAPH.LEFT):
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        p.paragraph_format.alignment = align

    def clear_and_add_runs(p, runs_spec, space_before=3, space_after=3, line_spacing=1.3, align=WD_ALIGN_PARAGRAPH.LEFT):
        clear_paragraph(p)
        format_paragraph(p, space_before=space_before, space_after=space_after, line_spacing=line_spacing, align=align)
        for text, size_pt, font_type, bold, color in runs_spec:
            r = p.add_run(text)
            format_academic_run(r, font_size_pt=size_pt, font_type=font_type, bold=bold, color=color)

    def set_cell(cell, text, bold=False, font_size_pt=9.0, font_type="body", bg_hex=None, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = cell.paragraphs[0]
        clear_paragraph(p)
        format_paragraph(p, space_before=2, space_after=2, line_spacing=1.15, align=align)
        run = p.add_run(text)
        format_academic_run(run, font_size_pt=font_size_pt, font_type=font_type, bold=bold, color=COLOR_BLACK)
        if bg_hex:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_hex}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

    # 1. Format Paragraphs (Academic Thesis Scale)
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue

        if "千问办公专家套件方案书" in txt:
            # 论文主标题：小二号 (18pt) 黑体 加粗 居中
            clear_and_add_runs(p, [("千问办公专家套件方案书", 18, "title", True, COLOR_HEADING)], space_before=12, space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif "专家名称：" in txt:
            # 论文作者/基本信息：五号 (10.5pt) 楷体 居中
            clear_and_add_runs(p, [("专家名称：", 10.5, "heading", True, COLOR_BLACK), ("多专家协同研发工作流 (multi-agent-flow)", 10.5, "meta", False, COLOR_BLACK)], space_before=3, space_after=2)
        elif "申报主体：" in txt:
            clear_and_add_runs(p, [("申报主体：", 10.5, "heading", True, COLOR_BLACK), ("任可", 10.5, "meta", False, COLOR_BLACK)], space_before=2, space_after=10)
        elif txt.startswith("一、") or txt.startswith("二、") or txt.startswith("三、") or txt.startswith("四、"):
            # 论文一级标题：四号 (14pt) 黑体 加粗
            clear_and_add_runs(p, [(txt, 14, "heading", True, COLOR_HEADING)], space_before=14, space_after=6)
        elif txt.startswith("2.1") or txt.startswith("2.2") or txt.startswith("2.3") or txt.startswith("4.1") or txt.startswith("4.2"):
            # 论文二级标题：小四号 (12pt) 黑体 加粗
            clear_and_add_runs(p, [(txt, 12, "heading", True, COLOR_HEADING)], space_before=10, space_after=4)
        elif "一句话定位：" in txt:
            clear_and_add_runs(p, [
                ("【一句话定位】 ", 10.5, "heading", True, COLOR_HEADING),
                ("面向敏捷研发团队与独立开发者的 8 位 AI 专家协同工作流套件，提供契约驱动的 5 层防错状态机、在制品并发控制与人类专属验收质量门禁。", 10.5, "body", False, COLOR_BLACK)
            ], space_before=4, space_after=4)
        elif "差异化壁垒：" in txt:
            barrier_text = (
                "1. 8 角色细粒度分工闭环：拆解为 PM 严经理、架构师钱架构、开发李开发、前端马前端、审查周审查、测试章测试、运维吕改特、文档李文通，避免单 Agent 自写自测自夸；\n"
                "2. 5 层防错状态机硬门禁：强类型 JSON 契约、在制品并发上限（WIP Limit）、防冲卡时序校验、缺陷原卡打回不拆单、Git Pre-commit 物理硬拦截；\n"
                "3. 人类专属【已验收】安全红线：严禁 AI Agent 自行标记已验收结项，必须由人类用户终审确认；\n"
                "4. 多客户端开箱即用：原生适配千问办公（Qoder）、Antigravity、Claude Code 等主流环境。"
            )
            clear_and_add_runs(p, [
                ("【差异化壁垒】\n", 10.5, "heading", True, COLOR_HEADING),
                (barrier_text, 10.5, "body", False, COLOR_BLACK)
            ], space_before=4, space_after=4)
        elif "目标用户一句话描述" in txt:
            clear_and_add_runs(p, [("使用千问办公/AI IDE 进行日常软件需求拆解、架构设计、编码测试与敏捷交付的技术主管、敏捷项目经理、全栈工程师及独立开发者。", 10.5, "body", False, COLOR_BLACK)], space_before=3, space_after=4)
        elif "核心痛点" in txt:
            pain_text = (
                "1. 单 Agent 角色混淆与盲目自夸：单个 AI 既当裁判又当选手，容易产生伪完工与缺陷逃逸；\n"
                "2. 多任务并发导致上下文爆炸与脏写冲突：缺乏在制品（WIP）并发上限控制，多任务同时修改引发代码污染；\n"
                "3. 过程失控与交付黑盒：缺少阶段门禁校验，AI 擅自跳过审查或测试直接提交，甚至擅自关闭任务。"
            )
            clear_and_add_runs(p, [(pain_text, 10.5, "body", False, COLOR_BLACK)], space_before=3, space_after=6)
        elif "如在专家套件中不封装 MCP" in txt:
            clear_and_add_runs(p, [("注：本套件采用本地原生轻量 CLI 引擎驱动，无需额外封装外部独立 MCP 连接器。", 9.5, "meta", False, COLOR_MUTED)], space_before=2, space_after=4)

    # 2. Table 0: Persona
    t0 = doc.tables[0]
    set_cell(t0.rows[0].cells[0], "要素", bold=True, font_size_pt=9.5, font_type="heading", bg_hex="F1F5F9", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(t0.rows[0].cells[1], "详细描述", bold=True, font_size_pt=9.5, font_type="heading", bg_hex="F1F5F9", align=WD_ALIGN_PARAGRAPH.CENTER)
    persona_data = [
        ("岗位", "技术主管 (Tech Lead) / 研发项目经理 (PM) / 全栈工程师 (Full-Stack Dev) / 独立开发者 (Indie Hacker)"),
        ("组织规模", "1~50 人的敏捷研发团队、数字化转型敏捷小组或独立开发工作室"),
        ("日常主要工作", "需求 WBS 拆解、架构设计与 ADR 决策、前后端业务编码、代码审查、自动化测试、缺陷修复、发布运维与工程文档沉淀"),
        ("常用工具", "千问办公、Qoder、VS Code、Git/GitHub、Docker、Postman、pytest/vitest"),
        ("核心 KPI", "研发迭代交付周期 (Lead Time)、缺陷逃逸率 (Defect Rate)、代码审查覆盖率 (CR Rate)、工程文档完备率"),
        ("成功指标", "AI 交付代码 0 悬挂、缺陷返工率降低 60% 以上、研发全流程各阶段状态与产物透明可溯")
    ]
    for row_idx, (k, v) in enumerate(persona_data, start=1):
        set_cell(t0.rows[row_idx].cells[0], k, bold=True, font_size_pt=9.0, font_type="body", bg_hex="F8FAFC")
        set_cell(t0.rows[row_idx].cells[1], v, bold=False, font_size_pt=9.0, font_type="body")

    # 3. Table 1: Scenes
    t1 = doc.tables[1]
    while len(t1.rows) > 0:
        t1._tbl.remove(t1.rows[-1]._tr)

    h_row = t1.add_row()
    h_titles = ["场景名称", "触发频次", "当前核心痛点", "输入数据", "期望输出产物", "是否有写操作/敏感数据"]
    for ci, ht in enumerate(h_titles):
        set_cell(h_row.cells[ci], ht, bold=True, font_size_pt=9.5, font_type="heading", bg_hex="F1F5F9", align=WD_ALIGN_PARAGRAPH.CENTER)

    scenes = [
        ("1. 项目全自动扫描与架构初始化", "项目创建 / 首次接入", "新项目代码结构复杂，AI 无法快速摸清架构上下文与规范", "项目根目录源码 / git 仓库", "生成 project_architecture.yaml 架构档案与 8 专家上下文", "是 (生成架构配置文件)"),
        ("2. 需求 WBS 细化与任务拆解", "高频 (每次迭代 / 故事开工)", "大需求模糊、边界不清，AI 容易漏做或过度实现", "自然语言需求 / 用户故事", "生成拆解后的原子任务卡与责任专家分配 (PM 严经理调度)", "是 (写入看板任务卡)"),
        ("3. 架构设计与 ADR 决策记录", "中频 (技术选型 / 结构重构)", "架构决策缺失沉淀，后期成为团队维护黑盒", "技术选型或重构方案", "标准化 ADR 架构决策记录与技术栈依赖更新", "是 (写入 docs 架构文档)"),
        ("4. 原子任务开发与编码交付", "极高频 (日常任务实施)", "编码与契约脱节，前后端协作接口不一致", "待办任务卡 + 架构上下文", "业务功能代码实现、关联单测用例与变更清单", "是 (编写项目源代码)"),
        ("5. 代码审查与门禁审计", "高频 (任务提测前)", "代码风格不统一、潜在漏洞未检出、AI 盲目自夸", "开发提交的代码 diff 与交付说明", "严格 Review 意见 (P0-P3 缺陷清单、重构建议、门禁判定)", "否 (只读审计与评估)"),
        ("6. 自动化测试与缺陷原卡打回", "高频 (提测验证阶段)", "缺少自动化验证，测试不通过却拆新单导致追溯困难", "待测代码与测试用例", "测试执行日志、覆盖率报告；缺陷时原卡打回挂载缺陷", "是 (执行测试与状态打回)"),
        ("7. 人类专属验收与结项上锁", "高频 (任务终态验收)", "AI 擅自结项跳过人工把关，导致不可控交付", "测试通过的任务卡与验收清单", "人类核验通过后，状态变更为【已验收】，任务卡结项上锁", "是 (写入终态已验收状态)"),
        ("8. 可视化多角色实时看板查询", "随时 (日常进度追踪)", "多人/多 Agent 协同进度不透明，无法直观掌握卡点", "无 / yy-flow kanban 指令", "浏览器实时渲染 3 列看板与 8 角色在制品状态", "否 (只读可视化查询)"),
        ("9. Git 提交前置硬门禁拦截", "高频 (每次 git commit)", "存在进行中或未验收任务卡时误提交代码污染主干", "本地 git commit 动作", "Pre-commit 钩子校验；存在非终态卡时物理中断拦截", "否 (执行安全门禁判定)"),
        ("10. 敏捷迭代复盘与文档沉淀", "中频 (迭代结项 / 发版)", "迭代完成无沉淀，技术负债与质量数据丢失", "本期已完成任务集与审计记录", "生成敏捷复盘总结报告与全量工程文档归档", "是 (写入复盘与沉淀文档)")
    ]
    for sc in scenes:
        row = t1.add_row()
        for col_idx, val in enumerate(sc):
            set_cell(row.cells[col_idx], val, bold=(col_idx==0), font_size_pt=9.0, font_type="body", bg_hex="F8FAFC" if col_idx==0 else None)

    # 4. Table 2: Skills
    t2 = doc.tables[2]
    while len(t2.rows) > 0:
        t2._tbl.remove(t2.rows[-1]._tr)

    h_row2 = t2.add_row()
    h_titles2 = ["Skill 名称", "触发场景", "输入参数/指令", "期望输出产物"]
    for ci, ht in enumerate(h_titles2):
        set_cell(h_row2.cells[ci], ht, bold=True, font_size_pt=9.5, font_type="heading", bg_hex="F1F5F9", align=WD_ALIGN_PARAGRAPH.CENTER)

    skills_data = [
        ("yy-flow", "全生命周期多专家协同工作流编排与状态推进", "自然语言指令 / /yy-flow", "8 位专家按序流转，推进任务自需求至验收"),
        ("yy-flow-start", "项目开工与自动化技术栈扫描初始化", "/yy-flow start [项目路径]", "完成架构扫描、生成多专家上下文与看板初始配置"),
        ("yy-flow-status", "多专家当前任务状态与在制品进度查询", "/yy-flow status", "输出终端结构化看板与各角色在制品卡片详情"),
        ("yy-flow-kanban", "启动本地可视化 Web 看板服务", "/yy-flow kanban [端口]", "启动 HTTP 实时看板服务并输出可视化浏览器看板"),
        ("yy-flow-metrics", "研发质量度量与合规数据分析", "/yy-flow metrics", "输出流转周期、打回率、WIP 并发与质量得分报表")
    ]
    for sk in skills_data:
        row = t2.add_row()
        for col_idx, val in enumerate(sk):
            set_cell(row.cells[col_idx], val, bold=(col_idx==0), font_size_pt=9.0, font_type="body", bg_hex="F8FAFC" if col_idx==0 else None)

    # 5. Table 3: MCP
    t3 = doc.tables[3]
    while len(t3.rows) > 0:
        t3._tbl.remove(t3.rows[-1]._tr)

    h_row3 = t3.add_row()
    h_titles3 = ["连接器 MCP 名称", "能力说明", "内置 Tool 工具清单"]
    for ci, ht in enumerate(h_titles3):
        set_cell(h_row3.cells[ci], ht, bold=True, font_size_pt=9.5, font_type="heading", bg_hex="F1F5F9", align=WD_ALIGN_PARAGRAPH.CENTER)

    mcp_row = t3.add_row()
    set_cell(mcp_row.cells[0], "多专家协同原生引擎\n(无外部独立 MCP)", bold=True, font_size_pt=9.0, font_type="body", bg_hex="F8FAFC")
    set_cell(mcp_row.cells[1], "本套件采用本地轻量 CLI 脚本引擎与 Git 钩子驱动，无需依赖外部独立 MCP 服务，零额外网络开销，支持完全离线与私有化安全运行。", font_size_pt=9.0, font_type="body")
    set_cell(mcp_row.cells[2], "内置 35 项流转、门禁、度量与看板核心脚本工具（含 transition_task, verify_git_gate, check_stage_gate 等）", font_size_pt=9.0, font_type="body")

    # Save
    doc.save(out_path)
    doc.save(download_copy)
    print(f"Academic Thesis Proposal saved to {out_path} and {download_copy}")

if __name__ == "__main__":
    generate_academic_proposal()
