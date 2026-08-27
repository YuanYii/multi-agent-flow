"""
DOCX Academic Thesis / Technical Report Styler (GB/T 7713)
提供标准学术论文排版引擎：宋体 (SimSun) + Times New Roman 正文、黑体 (SimHei) 标题、楷体 (KaiTi) 元数据，
并强制绑定 OpenXML eastAsia 东亚文字槽，杜绝中西文字体脱节与局部字体回退断层。
"""

import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_HEADING = RGBColor(17, 24, 39)
COLOR_MUTED = RGBColor(75, 85, 99)

def init_academic_document(top_margin_in=1.0, bottom_margin_in=1.0, left_margin_in=1.1, right_margin_in=1.1):
    """创建并初始化符合学术规范的 A4 Word 文档"""
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(top_margin_in)
        section.bottom_margin = Inches(bottom_margin_in)
        section.left_margin = Inches(left_margin_in)
        section.right_margin = Inches(right_margin_in)

    # 规范全局默认样式
    for s_name in ["Normal", "正文"]:
        if s_name in doc.styles:
            s = doc.styles[s_name]
            s.font.name = "Times New Roman"
            s.font.size = Pt(10.5)
            rFonts = s.element.rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), "宋体")
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:cs"), "Times New Roman")
    return doc

def clear_paragraph(p):
    """彻底清除段落中现存的所有 run，防止遗留未设置字体的空节点"""
    for r in list(p.runs):
        p._p.remove(r._r)

def format_academic_run(run, font_size_pt=10.5, font_type="body", bold=False, color=COLOR_BLACK):
    """
    格式化学术 Run：
    - font_type: 'title' / 'heading' -> 黑体 + Times New Roman
    - font_type: 'meta' / 'quote'   -> 楷体 + Times New Roman
    - font_type: 'body' / 'table'   -> 宋体 + Times New Roman
    """
    run.font.size = Pt(font_size_pt)
    run.bold = bold
    run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    if font_type in ["title", "heading"]:
        rFonts.set(qn("w:eastAsia"), "黑体")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")
    elif font_type == "meta":
        rFonts.set(qn("w:eastAsia"), "楷体")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")
    else:  # body / table
        rFonts.set(qn("w:eastAsia"), "宋体")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:hint"), "eastAsia")

def add_academic_title(doc, text):
    """添加学术论文大标题：小二号 (18pt) 黑体 加粗 居中"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    format_academic_run(run, font_size_pt=18, font_type="title", bold=True, color=COLOR_HEADING)
    return p

def add_academic_subtitle(doc, text):
    """添加作者/元数据：五号 (10.5pt) 楷体 居中"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    format_academic_run(run, font_size_pt=10.5, font_type="meta", bold=False, color=COLOR_MUTED)
    return p

def add_academic_h1(doc, text):
    """添加一级标题：四号 (14pt) 黑体 加粗"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    format_academic_run(run, font_size_pt=14, font_type="heading", bold=True, color=COLOR_HEADING)
    return p

def add_academic_h2(doc, text):
    """添加二级标题：小四号 (12pt) 黑体 加粗"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    format_academic_run(run, font_size_pt=12, font_type="heading", bold=True, color=COLOR_HEADING)
    return p

def add_academic_p(doc, text, bold=False):
    """添加正文段落：五号 (10.5pt) 宋体，1.3 倍行距"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    format_academic_run(run, font_size_pt=10.5, font_type="body", bold=bold, color=COLOR_BLACK)
    return p

def set_academic_cell(cell, text, bold=False, font_size_pt=9.0, font_type="body", bg_hex=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    """设置学术表格单元格文字：小五号 (9.0pt) 宋体/黑体"""
    p = cell.paragraphs[0]
    clear_paragraph(p)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.alignment = align
    run = p.add_run(text)
    format_academic_run(run, font_size_pt=font_size_pt, font_type=font_type, bold=bold, color=COLOR_BLACK)
    if bg_hex:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)


def setup_academic_page(doc, top_margin_in=1.0, bottom_margin_in=1.0,
                        left_margin_in=1.1, right_margin_in=1.1):
    """对已有 Document（新建或模板加载）应用 GB/T 7713 学术页边距与全局默认样式。

    与 init_academic_document 配套：模板场景用本函数补设页面，
    全新文档场景直接使用 init_academic_document 一步到位。
    """
    for section in doc.sections:
        section.top_margin = Inches(top_margin_in)      # 2.54 cm
        section.bottom_margin = Inches(bottom_margin_in)  # 2.54 cm
        section.left_margin = Inches(left_margin_in)    # 2.8 cm
        section.right_margin = Inches(right_margin_in)  # 2.8 cm

    for s_name in ["Normal", "正文"]:
        if s_name in doc.styles:
            s = doc.styles[s_name]
            s.font.name = "Times New Roman"
            s.font.size = Pt(10.5)
            rFonts = s.element.rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), "宋体")
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:cs"), "Times New Roman")
    return doc


def clear_and_add_runs(p, runs_spec, space_before=3, space_after=3,
                       line_spacing=1.3, align=WD_ALIGN_PARAGRAPH.LEFT):
    """清空既有段落并按规格批量写入学术 Run（多行混排场景统一入口）。

    runs_spec: [(text, size_pt, font_type, bold, color), ...]
    """
    clear_paragraph(p)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.alignment = align
    for text, size_pt, font_type, bold, color in runs_spec:
        r = p.add_run(text)
        format_academic_run(r, font_size_pt=size_pt, font_type=font_type, bold=bold, color=color)
    return p
